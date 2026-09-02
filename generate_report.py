#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Staj raporu üretici (python-docx).
StajRaporuDuzeni.pdf içindeki biçim kurallarına uyar:
  - Arial 12 punto
  - Kenar boşlukları: üst/sol 4.0 cm, alt/sağ 2.5 cm
  - 1.5 satır aralığı
  - Ana başlık: BÜYÜK HARF + koyu
  - Alt başlık (<=3 seviye): Her Sözcüğün İlk Harfi Büyük + koyu
  - Numaralandırma: 1. / 1.1 / 2.1.1 / a)
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn


def _set_arial(run, size=12):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Arial")
    rfonts.set(qn("w:cs"), "Arial")


def yeni_belge() -> Document:
    doc = Document()

    for bolum in doc.sections:
        bolum.top_margin = Cm(4.0)
        bolum.left_margin = Cm(4.0)
        bolum.bottom_margin = Cm(2.5)
        bolum.right_margin = Cm(2.5)

    stil = doc.styles["Normal"]
    font = stil.font
    font.name = "Arial"
    font.size = Pt(12)
    rpr = stil.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Arial")
    rfonts.set(qn("w:cs"), "Arial")

    pf = stil.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    return doc


def _baslik(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.bold = True
    _set_arial(run)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.space_before = Pt(6)
    return p


def ana_baslik(doc, text):
    _baslik(doc, text)


def alt_baslik(doc, text):
    _baslik(doc, text)


def paragraf(doc, text, girinti_cm=0.0):
    p = doc.add_paragraph(text)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if girinti_cm:
        pf.left_indent = Cm(girinti_cm)
    for run in p.runs:
        _set_arial(run)
    return p


def madde(doc, text, girinti_cm=1.0, isaret="•"):
    p = doc.add_paragraph(f"{isaret} {text}")
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent = Cm(girinti_cm)
    for run in p.runs:
        _set_arial(run)
    return p


def kod_parca(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    return p


def ara(doc):
    doc.add_paragraph("")


def kapak(doc):
    for _ in range(2):
        ara(doc)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("STAJ RAPORU")
    run.bold = True
    _set_arial(run, size=24)

    for _ in range(3):
        ara(doc)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Kod Tabanına Özel Geri Getirmeli Üretim (RAG) Tabanlı\nYerel Dil Modeli Ajan Sistemi")
    run.bold = True
    _set_arial(run, size=16)

    for _ in range(4):
        ara(doc)

    bilgiler = [
        ("Ad Soyad", "......................................."),
        ("Sınıf", "......................................."),
        ("Numara", "......................................."),
        ("Stajın Konusu", "Yazılım"),
        ("Başlama Tarihi", ".... / .... / ......."),
        ("Bitiş Tarihi", ".... / .... / ......."),
        ("Staj Süresi", ".... / ...."),
        ("Staj Yapılan Kurum", "......................................."),
        ("Kurum Adresi", "......................................."),
    ]
    for etiket, deger in bilgiler:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(f"{etiket}: ")
        run.bold = True
        _set_arial(run)
        run2 = p.add_run(deger)
        _set_arial(run2)

    doc.add_page_break()


def icindekiler(doc):
    ana_baslik(doc, "İÇİNDEKİLER")
    ara(doc)
    satirlar = [
        ("1. GİRİŞ", "1"),
        ("2. KURUM TANITIMI", "1"),
        ("3. STAJ SÜRECİ", "2"),
        ("3.1. Proje Konusu ve Genel Mimari", "2"),
        ("3.2. Kullanılan Teknolojiler", "3"),
        ("3.3. AST Tabanlı Kod Parçalama (splitters)", "4"),
        ("3.4. Qdrant Vektör Veritabanı İndeksleme Hattı (qdrant_pipeline)", "5"),
        ("3.5. LangChain / LangGraph Ajan Mimarisi (local_llm ve interface)", "6"),
        ("3.5.1. Tanımlanan Araçlar (Tools)", "6"),
        ("3.5.2. Sistem Komutlandırması (promptlar)", "7"),
        ("3.5.3. Bellek ve Denetim Noktası (MemorySaver)", "7"),
        ("3.6. Streamlit Web Arayüzü (interface/app.py)", "8"),
        ("3.6.1. Oturum Yönetimi ve Diskte Süreklilik", "8"),
        ("3.6.2. Dosya Paneli ve Tam Ekran Görüntüleme", "9"),
        ("3.7. Uzun Vadeli Bellek: Compaction Mekanizması", "9"),
        ("3.8. Token Takibi (token_tracker)", "10"),
        ("4. SONUÇ", "11"),
        ("5. KAYNAKLAR", "12"),
    ]
    for basl, sayfa in satirlar:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        tab = "\t" + "." * 4
        run = p.add_run(f"{basl}{tab}{sayfa}")
        run.bold = True
        _set_arial(run)
    doc.add_page_break()


def giris(doc):
    ana_baslik(doc, "1. GİRİŞ")
    ara(doc)

    paragraf(doc, "Bu staj çalışmasının konusu, büyük dil modeli (LLM) tabanlı bir yapay zekâ "
                  "asistanının, kaynak kod tabanlarını doğal dil üzerinden sorgulanabilir hâle "
                  "getiren bir Geri Getirmeli Üretim (Retrieval-Augmented Generation, RAG) "
                  "sisteminin tasarlanması ve geliştirilmesidir.")

    paragraf(doc, "Çalışmanın amacı; Java, Python ve Go dillerinde yazılmış kod "
                  "deposundan anlamlı parçalar çıkararak bunları bir vektör veritabanında "
                  "indekslemek ve kullanıcının sorularına yalnızca bu kod bağlamından "
                  "faydalanan yerel bir ajan üzerinden doğru ve denetlenebilir yanıtlar "
                  "üretmektir. Sistem, çıktılarını oturum bazında diske kaydedecek biçimde "
                  "tasarlanmış; hem komut satırı hem de web tabanlı iki arayüzle "
                  "sunulmuştur. Bütün çıkarım ve gömme (embedding) işlemleri yerel "
                  "olarak çalıştırılan Ollama üzerinden yürütüldüğünden harici bir "
                  "API anahtarı gerektirmemektedir.")

    paragraf(doc, "Staj süresince yürütülen başlıca çalışmalar şunlardır: tree-sitter "
                  "soyut söz dizimi ağacı (AST) ayrıştırıcısıyla fonksiyon/metot "
                  "bazında kod parçalama; her parçayı zengin üst verilerle Qdrant "
                  "koleksiyonunda indeksleme; LangChain ve LangGraph çatısı altında, "
                  "araç çağrısı yapabilen durumlu bir ajan kurma; sohbet geçmişinin "
                  "diske kalıcı yazılması ve yeniden yüklenmesi; bağlam penceresini "
                  "yönetmek için katmanlı bir bellek sıkıştırma (compaction) mekanizması "
                  "geliştirme; ve son olarak çoklu oturum destekli bir Streamlit web "
                  "arayüzü implementasyonu.")

    paragraf(doc, "Çalışmanın sonucunda, kaynak kod çağında anlamsal arama yapabilen, "
                  "yanıtlarını oturum klasörüne dosya olarak kaydedebilen, uzun "
                  "süreli belleği sıkıştırma ile yönetebilen ve tamamen yerel "
                  "altyapı üzerinde çalışabilen üretim düzeyinde bir RAG sistemi "
                  "elde edilmiştir.")


def kurum(doc):
    ana_baslik(doc, "2. KURUM TANITIMI")
    ara(doc)

    paragraf(doc, "Stajın yapıldığı kurumun adı, adresi, temel uğraş alanı, kısa "
                  "tarihçesi ve örgüt yapısı bilgileri aşağıda bırakılmıştır; "
                  "ilgili alanlar staj sonunda kurum tarafından doldurulacaktır.")

    paragraf(doc, "Kurum Adı: .......................................................")
    paragraf(doc, "Adres: .......................................................")
    paragraf(doc, "Temel Uğraş Alanı: .......................................................")

    paragraf(doc, "Staj yapılan birimde kullanılan özel donanım ve yazılımlar ile "
                  "birimde yürütülen uygulamalara ilişkin bilgiler kurum tarafından "
                  "eklenecektir.")


def staj_sureci(doc):
    ana_baslik(doc, "3. STAJ SÜRECİ")
    ara(doc)

    # 3.1
    alt_baslik(doc, "3.1. Proje Konusu ve Genel Mimari")
    paragraf(doc, "Geliştirilen sistem; kaynak kod repolarını ayrıştırıp gömme "
                  "vektörlerine dönüştüren bir indeksleme hattı (splitters ve "
                  "qdrant_pipeline), LangGraph tabanlı durumlu bir ajan "
                  "(local_llm ve interface), yapılandırumlı sistem komutları "
                  "(prompts) ve token takibi ile bellek sıkıştırması "
                  "(token_tracker, compaction.md) bileşenlerinden oluşur. "
                  "Sistemin iki ön yüzü vardır: başsız komut satırı arayüzü "
                  "(local_llm/agent.py) ve üretim düzeyindeki Streamlit web "
                  "uygulaması (interface/app.py). Her ikisi de aynı araçları, "
                  "aynı sistem komutunu ve aymı yerel modeli kullanır.")

    paragraf(doc, "Üç temel yürütme girişi (entry point) tanımlanmıştır:")
    madde(doc, "streamlit run interface/app.py — üretim kullanıcı arayüzü; "
               "çoklu oturum, dosya paneli ve kalıcı sohbet geçmişi sağlar.")
    madde(doc, "python -m local_llm.agent [-id SESSION_ID] — uçbirim "
               "sohbet arayüzü; sunucu tarafı testler için sade varyant.")
    madde(doc, "python -m qdrant_pipeline.qdrant_main — /tmp/Repos "
               "kök dizinindeki tüm depoları tarayıp codebase "
               "koleksiyonunu sıfırdan kurar.")

    paragraf(doc, "Genel akış; kod deposunun AST üzerinden metot/fonksiyon "
                  "tanesi ayrıştırılması, her parçanın qwen3-embedding:0.6b "
                  "modeliyle 1024 boyutlu vektöre gömülmesi, vektörlerin "
                  "kosinüs benzerliği ile Qdrant içinde saklanması ve "
                  "kullanıcı sorusunun aynı modelle gömülüp en yakın "
                  "komşu aranması biçiminde çalışır.")

    # 3.2
    alt_baslik(doc, "3.2. Kullanılan Teknolojiler")
    paragraf(doc, "Proje kapsamında kullanılan başlıca teknolojiler ve "
                  "kullanım amaçları aşağıdaki tabloda özetlenmiştir.")
    tablo_basliklari = ["Kategori", "Teknoloji", "Kullanım Yeri"]
    satir = [
        ["Vektör Veritabanı", "Qdrant (qdrant_client)", "qdrant_pipeline/qdrant_main.py"],
        ["LLM Çalışma Zamanı", "Ollama + langchain_ollama.ChatOllama", "local_llm/agent.py, interface/app.py"],
        ["Gömme Modeli", "qwen3-embedding:0.6b (1024 boyut)", "qdrant_main.py ve search_codebase araçları"],
        ["Sohbet Modelleri", "glm-5.2:cloud, ornith:35b (başlık üretimi)", "interface/app.py"],
        ["AST Ayrıştırma", "tree-sitter + Java/Python/Go dilleri", "splitters/text_splitters.py"],
        ["Ajan Çatısı", "LangChain create_agent, langchain_core", "tüm entegrasyon katmanları"],
        ["Durumlu Ajan", "LangGraph MemorySaver + update_state", "agent.py:242, app.py:125"],
        ["Web Arayüzü", "Streamlit (st.dialog, st.popover, st.status)", "interface/app.py"],
        ["Yapılandırma", "python-dotenv (.env)", "tüm modüller"],
        ["Eş Zamanlılık", "threading.Thread (arka plan başlık)", "interface/app.py"],
        ["HTTP", "requests (Ollama gömme uç noktası)", "qdrant_main.py, search_codebase"],
    ]
    tablo = doc.add_table(rows=1, cols=3)
    tablo.style = "Light Grid Accent 1"
    hdr = tablo.rows[0].cells
    for i, h in enumerate(tablo_basliklari):
        hdr[i].text = ""
        pr = hdr[i].paragraphs[0]
        run = pr.add_run(h)
        run.bold = True
        _set_arial(run)
    for s in satir:
        row = tablo.add_row().cells
        for i, val in enumerate(s):
            row[i].text = ""
            pr = row[i].paragraphs[0]
            run = pr.add_run(val)
            _set_arial(run)
    ara(doc)

    # 3.3
    alt_baslik(doc, "3.3. AST Tabanlı Kod Parçalama (splitters)")
    paragraf(doc, "RAG sistemlerinde ayrıştırma (chunking) kalitesi doğrudan "
                  "geri getirme kalitesini belirler. Bu projede karakter veya "
                  "satır sayısına bağlı kaba ayrıştırma yerine, tree-sitter ile "
                  "kaynak kodun soyut söz dizimi ağacı (AST) gezilerek her "
                  "fonksiyon/metot tek başına yetkinen bir anlamsal birim olarak "
                  "çıkarılmıştır (splitters/text_splitters.py:11-15).")

    paragraf(doc, "PARSERS sözlüğü .java, .py ve .go uzantılarını ilgili dil "
                  "gramerlerinden oluşturulmuş tree-sitter Parser örneklerine "
                  "eşler. get_package_or_module fonksiyonu kök düğümden Java "
                  "package_declaration ve Go package_clause ifadelerini okuyarak "
                  "ad alanı bilgisini çıkarır (text_splitters.py:18-37).")

    paragraf(doc, "visit üreteci (generator) AST'yi özyinelemeli biçimde gezer; "
                  "class_definition / class_declaration bağlamını izler ve "
                  "method_declaration veya function_definition düğümlerini "
                  "yield eder (text_splitters.py:40-83). Her alt gezintiye "
                  "current_context.copy() verilerek kapsam korunur "
                  "(text_splitters.py:83).")

    paragraf(doc, "extract_node_data fonksiyonu çıkarılan her kod bloğu için "
                  "üç temel bileşen hazırlar (text_splitters.py:86-115):")
    madde(doc, "code: Düğümün ham kaynak metni (node.text.decode()).", girinti_cm=1.0)
    madde(doc, "comment: Hemen önce gelen kardeş (prev_named_sibling) düğüm "
               "bir block_comment, comment veya expression_statement ise bunu "
               "yakalar; Javadoc/GoDoc/Python docstring ifadelerinin korunmasını "
               "sağlar (text_splitters.py:90-95).")
    madde(doc, "embedding_text: comment + '\\n\\n' + kod biçiminde "
               "birleştirilir; gömme modeli anlam yönünden zenginleştirilir "
               "(text_splitters.py:95).")

    paragraf(doc, "Üst veri (metadata) sözlüğü project, language, "
                  "package_or_module, class, type (method/function), name, "
                  "start_line, end_line ve file_path alanlarını taşır "
                  "(text_splitters.py:97-107). start_line ve end_line "
                  "sayesinde yanıtlar kesin dosya/satır referanslarıyla "
                  "verilebilmektedir.")

    paragraf(doc, "parse_all_repositories kök dizini (/tmp/Repos) özyinelemeli "
                  "gezer; her dosya için proje adını ilk yol bileşeninden, "
                  "dil adını lang_map'ten türetir ve dosyayı bayt olarak okuyup "
                  "ayrıştırır (text_splitters.py:118-154). Üreteç tasarımı "
                  "büyük kod tabanlarında bellek verimliliği sağlar. Aşağıda "
                  "temel ayrıştırıcı yapı verilmiştir:")

    kod_parca(doc,
        "PARSERS = {\n"
        "    \".java\": Parser(Language(tree_sitter_java.language())),\n"
        "    \".py\":   Parser(Language(tree_sitter_python.language())),\n"
        "    \".go\":   Parser(Language(tree_sitter_go.language())),\n"
        "}\n"
        "...\n"
        "if prev and prev.type in [\"block_comment\", \"comment\", \"expression_statement\"]:\n"
        "    comment = prev.text.decode().strip()\n"
        "embedding_text = f\"{comment}\\n\\n{code}\" if comment else code")
    ara(doc)

    # 3.4
    alt_baslik(doc, "3.4. Qdrant Vektör Veritabanı İndeksleme Hattı (qdrant_pipeline)")
    paragraf(doc, "qdrant_pipeline/qdrant_main.py modülü indeksleme katmanıdır. "
                  "Yerel Qdrant sunucusuna (localhost:6333) bağlanır; VECTOR_SIZE = 1024 "
                  "(qwen3-embedding modeli boyutu) ve COLLECTION_NAME = "
                  "\"codebase\" olarak tanımlanır (qdrant_main.py:14-18).")

    paragraf(doc, "Betiğin __main__ bloğu (qdrant_main.py:25-80) idempotent "
                  "(tekrar çalıştırılabilir) biçimde tasarlanmıştır: codebase "
                  "koleksiyonu zaten varsa önce silinir, ardından kosinüs "
                  "uzaklığı (Distance.COSINE) ile yeni koleksiyon oluşturulur "
                  "(qdrant_main.py:27-37). parse_all_repositories üretecinin "
                  "her parçası için embedding_text, yerel Ollama gömme uç "
                  "noktasına (EMBED_URL) model \"qwen3-embedding:0.6b\" "
                  "ile POST edilir; dönen ilk gömme vektörü çekilir "
                  "(qdrant_main.py:48-56).")

    paragraf(doc, "Her Qdrant noktası (point) tam kendini tanımlar: payload, "
                  "chunk üst verilerine ek olarak code ve comment alanlarını "
                  "içerir (qdrant_main.py:59-63). Noktalar 1'den başlayan "
                  "artan tamsayı kimlikleriyle upsert edilir; her ekleme "
                  "konsola dil, dosya yolu ve metot adı olarak loglanır "
                  "(qdrant_main.py:66-77). Bu sayede her geri getirme "
                  "sonucunda ham kod, üst veri ve skor birlikte alınır.")

    # 3.5
    alt_baslik(doc, "3.5. LangChain / LangGraph Ajan Mimarisi (local_llm ve interface)")
    paragraf(doc, "Ajan, langchain.agents.create_agent ile LangGraph "
                  "MemorySaver denetim noktası (checkpointer) kullanılarak "
                  "kurulur (agent.py:242-249; app.py:127-132). thread_id "
                  "yapılandırması her sohbet oturumuna bağımsız bellek "
                  "verir; araç çağrıları ve araç yanıtları bu durum içinde "
                  "korunur. Ajan, glm-5.2:cloud modelini sicaklik 0.4 ve "
                  "geniş bağlam penceresiyle kullanır (CLI num_ctx=262144, "
                  "web num_ctx=1000000).")

    alt_baslik(doc, "3.5.1. Tanımlanan Araçlar (Tools)")
    paragraf(doc, "@tool dekoratörüyle tanımlanan üç araç vardır (agent.py:130-216; "
                  "app.py:48-110):")
    madde(doc, "search_codebase(new_question, limit): Soruyu aynı Ollama "
               "gömme modeliyle vektöre çevirir, Qdrant'ta en yakın komşu "
               "araması yapar; score > 0.4 eşiğinden büyük vuruşları "
               "payload + skor biçiminde birleştirip döndürür "
               "(agent.py:189-213).")
    madde(doc, "update_query(user_question): Soruyu iyileştirmek için yer "
               "tutan, boş bir kısım döndüren bir araçtır; sistem komutu bu "
               "aracı kullanarak sorgu yeniden yazımını yönlendirir "
               "(agent.py:182-186).")
    madde(doc, "create_file: Ajanın yanıtını diske oturum klasörüne kaydeder. "
               "Web varyantı RunnableConfig'ten thread_id'yi okur ve "
               "Path(file_name).name ile dosya adını temizleyerek yol "
               "gezinme (path traversal) saldırılarını engeller (app.py:48-81).")

    alt_baslik(doc, "3.5.2. Sistem Komutlandırması (promptlar)")
    paragraf(doc, "prompts/prompts2.py dosyası, ajanın davranışını belirleyen "
                  "yapılandırumlu sistem komutunu (SYSTEM_PROMPT) içerir. "
                  "Komut; Role and Persona, Scope and Limitations, Context "
                  "& Search Workflow ve File Generation & Output Management "
                  "olmak üzere Markdown başlıklarına ayrılır "
                  "(prompts2.py:2-30). Sistem strict biçimde yalnız yazılım "
                  "konulu sorulara yanıt verir; yazılım dışı sorulara "
                  "sabit Türkçe bir ret metni döndürülür. Bağlam eksikse "
                  "search_codebase çağrılır; ilk limit 3'tür; sonuç yoksa "
                  "project / language / file_path alanları genişletilir. "
                  "Her yanıt .java/.py/.go veya .md dosyası olarak "
                  "kaydedilir ve nihai sohbet yanıtinda yalniz dosya adi "
                  "ile durum bilgisi verilir. prompts/prompts.py dosyasi "
                  "ise tamamen Türkçe ilk sürümü içerir.")

    alt_baslik(doc, "3.5.3. Bellek ve Denetim Noktası (MemorySaver)")
    paragraf(doc, "LangGraph MemorySaver, her thread_id için bağımsız bir "
                  "durum saklar. session_id, uuid4 ile üretilir; -id "
                  "argümanı verilirse önceki oturum diskten yüklenir ve "
                  "continue_session CONTEXT*.md dosyalarını okuyarak "
                  "bir SystemMessage olarak duruma enjekte eder "
                  "(agent.py:46-67). create_agent'in MemorySaver checkpointer "
                  "ile kullanımı, LangGraph'in durum azaltıcı (reducer) "
                  "olarak mesaj listesini saklaması sayesinde gerçekleşir; bu olmadan "
                  "RemoveMessage ile belirli mesajları silmek ve yeni "
                  "SystemMessage enjekte etmek olanaksız olurdu.")

    # 3.6
    alt_baslik(doc, "3.6. Streamlit Web Arayüzü (interface/app.py)")
    paragraf(doc, "interface/app.py (~622 satır), projenin öne çıkan "
                  "üretim girişidir. @st.cache_resource dolayısıyla ajan "
                  "ve LLM örnekleri Streamlit çalışması boyunca bir kez "
                  "oluşturulur (app.py:113-136). Arayüz, 0.35 oranla yan "
                  "dosya panelini görecek iki kolon (col_chat ve col_files) "
                  "biçiminde düzenlenmiştir (app.py:375-378).")

    alt_baslik(doc, "3.6.1. Oturum Yönetimi ve Diskte Süreklilik")
    paragraf(doc, "Sol kenar çubuğu \"Sohbet Geçmişi\" başlığı "
                  "altında yeni sohbet başlatma ve kayıtlı oturumları "
                  "listeleme/silme işlevlerini sunar (app.py:331-373). "
                  "Her oturum SESSION_PATH altında kendi klasörüne sahiptir; "
                  "UUID tabanlı kimlik thread_id olarak kullanılır.")

    paragraf(doc, "save_messages_to_disk ve load_messages_from_disk "
                  "fonksiyonları LangGraph mesajlarını tip bilgileri "
                  "korunarak history.json olarak serileştirir "
                  "(app.py:196-270). Özellikle AIMessage üzerinde tool_calls "
                  "ve ToolMessage üzerinde tool_call_id alanlarının saklanması "
                  "kritiktir; yanlış geri yükleme LangGraph'in araç "
                  "çağrısı korelasyonunu bozardı. Bu sayede sohbet "
                  "Streamlit yeniden çalıştırılmalarında ve uygulama "
                  "yeniden başlatmalarında bile korunur.")

    alt_baslik(doc, "3.6.2. Dosya Paneli ve Tam Ekran Görüntüleme")
    paragraf(doc, "Sağ panel, \"Session Dosya Paneli\" "
                  "genişletici altında oturumda üretilen dosyaları "
                  "listeler (app.py:530-622). Dosya türüne göre simge "
                  "(Python, txt/md, json/yaml, diğer) atanır; "
                  "popover üzerinden dosya seçimi yapılır. Markdown/metin "
                  "dosyaları için hazır görünüm ve ham metin sekmeleri, "
                  "kod dosyaları için satır numaralı sözdizimi vurgulu "
                  "st.code görünümü sunulur. @st.dialog ile tam ekran "
                  "modal, içerik/ham düzen sekmeleri ve indirme "
                  "düğmesi içerir.")

    # 3.7
    alt_baslik(doc, "3.7. Uzun Vadeli Bellek: Compaction Mekanizması")
    paragraf(doc, "Uzun sohbetlerde bağlam penceresinin tükenmesini önlemek "
                  "için, el yapımı (hand-crafted) bir yapılandırumlu özet "
                  "kullanılır. compaction.md şablonu, LLM'den sabit "
                  "Markdown bölümleri üretmesini ister: Objective, "
                  "Important Details, Work State (Completed / Active / "
                  "Blocked), Next Move ve Relevant Files (compaction.md).")

    paragraf(doc, "run_compaction fonksiyonu (CLI: agent.py:75-126; "
                  "web: app.py:290-329) şu adımları izler:")
    madde(doc, "compaction.md şablonunu ve tam sohbet geçmişini birleştirip "
               "LLM'i çağırır (agent.py:77-85).")
    madde(doc, "Özet CONTEXT{N}.md olarak oturum klasörüne kaydedilir; N, "
               "klasördeki mevcut context sayısıdır (agent.py:91-97).")
    madde(doc, "Eski mesajlar LangGraph'tan RemoveMessage ile silinir; "
               "CLI tümünü siler, web ise son KEEP_LAST_N = 6 mesajı korur "
               "(agent.py:103-106; app.py:316-322).")
    madde(doc, "Tüm CONTEXT*.md dosyaları birleştirilip yeni bir SystemMessage "
               "olarak duruma enjekte edilir (agent.py:118-123).")

    paragraf(doc, "Sıkıştırma iki yolla tetiklenir: kullanıcı /compact "
                  "yazdığında veya token takibi belirli bir eşiği "
                  "aştığında otomatik. Eşikler, bağlam penceresine göre "
                  "ayrı seçilmiştir: CLI'da 60 000 (agent.py:311), web "
                  "arayüzünde 750 000 (app.py:496).")

    # 3.8
    alt_baslik(doc, "3.8. Token Takibi (token_tracker)")
    paragraf(doc, "token_tracker/base_callback_handler.py dosyasındaki "
                  "TokenTrackerHandler, langchain_core.callbacks."
                  "BaseCallbackHandler alt sınıfıdır (base_callback_handler.py:1-17). "
                  "on_llm_end geri çağrısında, Ollama yanıt metadata'sındaki "
                  "prompt_eval_count ve eval_count alanlarını okuyarak "
                  "toplam giriş ve çıkış token sayisini biriktirir. Bu "
                  "değer hem gözlemlenebilirlik hem de sıkıştırma eşiği "
                  "tetikleyici olarak kullanılır ve her agent adimindan "
                  "sonra konsola yanit bilgilere yansitilir "
                  "(agent.py:303-306).")

    paragraf(doc, "Web arayüzünde, ikincil bir model ornith:35b "
                  "(create_ornith, app.py:140-148), kullanıcının ilk "
                  "mesajından 3-5 kelimelik Türkçe sohbet başlığı üretir. "
                  "Bu çağrı, arayüzün beklememesi için bir daemon "
                  "threading.Thread üzerinde arka planda yapılır "
                  "(app.py:451-462).")


def sonuc(doc):
    ana_baslik(doc, "4. SONUÇ")
    ara(doc)

    paragraf(doc, "Bu staj, kaynak koda özel tam özellikli bir RAG "
                  "sisteminin uçtan uca tasarlanması ve geliştirilmesi "
                  "bağlamında önemli bir mühendislik deneyimi sağlamıştır. "
                  "Yapay zekâ odaklı proje geliştirme sürecinin planlamadan "
                  "uygulamaya tüm aşamaları bire bir deneyimlenmiş; vektör "
                  "veritabanı, AST ayrıştırma, ajan çatıları ve sohbet "
                  "arayüzü teknikleri bir arada kullanılmıştır.")

    paragraf(doc, "Stajın kazandırdığı başlıca bilgi ve deneyimler şunlardır:")
    madde(doc, "RAG sistemlerinde ayrıştırma biriminin kalitesinin geri "
               "getirme üzerine doğrudan etkisini görmek; AST tabanlı metot "
               "bazlı ayırmanın kaba satır/karakter ayırmalarına göre "
               "anlamsal isabeti belirgin biçimde artırdığını gözlemlemek.")
    madde(doc, "LangGraph'in durum azaltıcı mesaj modelinin, RemoveMessage "
               "ve update_state işlemleriyle nasıl gerçek bellek yönetimi "
               "sağladığını öğrenmek; bu sayede katmanlı compaction "
               "(sıkıştırma) mekanizması kurulması.")
    madde(doc, "Hibrit geri getirme potansiyeli için üst veri zenginliğinin "
               "(project, language, package_or_module, class, "
               "start_line, end_line, file_path) ne kadar önemli "
               "olduğunu görmek; vuruşların kesin dosya/satır "
               "referanslarıyla yanıtlanması.")
    madde(doc, "Streamlit gibi hızlı prototiplenebilir bir arayüz çatısıyla "
               "üretim düzeyinde oturum yönetimi, disk kalicılığı ve "
               "modal/popover diyaloğu kurgulamak.")
    madde(doc, "Yerel öncelikli (local-first) çıkarım yaklaşımının harici "
               "API olmadan Ollama ve Qdrant ile sağlanabileceğini "
               "gözlemlemek.")

    paragraf(doc, "Görev alınan projenin tamamlanan ve katkıda bulunulan kısımları:")
    madde(doc, "splitters/text_splitters.py içinde AST tabanlı dil farkındalıklı "
               "kod parçalama ve comment koruma mantığının kuruluşu.")
    madde(doc, "qdrant_pipeline/qdrant_main.py içinde idempotent indeks "
               "hattı ve zengin üstdolu payload tasarımı.")
    madde(doc, "Araç çağrısına izin veren LangGraph stateful ajan "
               "kurulması ve sistem komutunun iki sürüm üzerinden "
               "iyileştirilmesi.")
    madde(doc, "Uzun süreli bellek için yapılandınumlu compaction.md "
               "şablonu ve katmanlı CONTEXT dosya yönetimi.")
    madde(doc, "Web arayüzünde tip koruyan history.json serileştirilmesi, "
               "arayüz güvenliği (Path temizleme) ve arka plan başlık "
               "üretimi.")

    paragraf(doc, "Süreçte gözlemlenen başlıca sorunlar ve yorumlar:")
    madde(doc, "Saf vektör kosinüs aramasının çok projede ad çakışmaları "
               "nedeniyle yanlış vuruş getirebildiği; score > 0.4 eşiğinin "
               "her kod tabanında optimal olmadığı - üst veri bazlı "
               "filtrelerle hibrit arama çözüm olabilir.")
    madde(doc, "Komut satırı varyantında Python AST mantığı fonksiyon "
               "adını class bağlamı olarak yeniden atar; bu "
               "davranışın sınıf kaybına yol açtığı görüldü "
               "(text_splitters.py:64-65, 72-78).")
    madde(doc, "Qdrant toplu gömme istekleri yerine tek tek upsert "
               "kullanmaktadır; büyük depolarda indeksleme süresi "
               "uzamaktadır.")
    madde(doc, "Egg-info kaynak listesinin local_llm/core.py ve "
               "splitters/tree_splitter.py dosyalarına atıfta bulunması "
               "ancak bu dosyaların var olmaması, paketlenmenin "
               "tutarsız olduğunu göstermektedir.")

    alt_baslik(doc, "4.1. Teknik Öneriler")
    madde(doc, "Qdrant'ta vektör + alan filtresi (metadata filtering) "
               "desteği kullanılarak hibrit geri getirme yapılandırılmalıdır.")
    madde(doc, "İndeksleme sırasında batch gömme ve upsert toplu işleme "
               "kullanılmalıdır.")
    madde(doc, "Python dilinde fonksiyonların class bağlamı, gerçek "
               "dekoratör bağlamından ayrılmadan korunmalıdır.")
    madde(doc, "İndeksleme artımlı, yalnızca değişen dosyaları yeniden "
               "gömme olacak biçimde desteklenmelidir.")
    madde(doc, "Prompt güvenliği için sistem komutunun girilen dosya adı "
               "doğrulaması create_file aracında CLI tarafında da "
               "uygulanmalıdır.")


def kaynaklar(doc):
    ana_baslik(doc, "5. KAYNAKLAR")
    ara(doc)
    kaynak_listesi = [
        "LangChain Belgeleri, https://python.langchain.com",
        "LangGraph Belgeleri, https://langchain-ai.github.io/langgraph",
        "Qdrant Documentation, https://qdrant.tech/documentation",
        "Ollama, https://ollama.com",
        "tree-sitter, https://tree-sitter.github.io/tree-sitter",
        "Streamlit Documentation, https://docs.streamlit.io",
        "python-docx Documentation, https://python-docx.readthedocs.io",
        "LangChain Ollama Entegrasyonu, langchain-ollama",
        "Embedding Modeli: qwen3-embedding:0.6b",
        "Sohbet Modeli: glm-5.2:cloud; Başlık Modeli: ornith:35b",
    ]
    for i, k in enumerate(kaynak_listesi, start=1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-1.0)
        run = p.add_run(f"[{i}] ")
        run.bold = True
        _set_arial(run)
        run2 = p.add_run(k)
        _set_arial(run2)


def main():
    doc = yeni_belge()
    kapak(doc)
    icindekiler(doc)
    giris(doc)
    kurum(doc)
    staj_sureci(doc)
    sonuc(doc)
    kaynaklar(doc)
    out = "/Users/efeemirhandogan/Documents/RagLLM/rag_llm/Staj_Raporu.docx"
    doc.save(out)
    print(f"Kaydedildi: {out}")


if __name__ == "__main__":
    main()